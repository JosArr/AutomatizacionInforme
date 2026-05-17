# src/document/template_handler.py
"""Manejador de templates Word"""

from pathlib import Path
from typing import Optional, List
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.logging_config.logger import get_logger


logger = get_logger()


class TemplateHandler:
    """Manejador de templates Word"""

    def __init__(self, template_path: Path = None):
        if template_path is None:
            template_path = Path(__file__).parent.parent.parent / "templates" / "INFORME_TEMPLATE.docx"

        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Template no encontrado: {self.template_path}")

        self.document = None

    def load_template(self) -> Document:
        """Cargar template desde archivo"""
        try:
            self.document = Document(self.template_path)
            logger.info(f"Template cargado: {self.template_path}")
            return self.document
        except Exception as e:
            logger.error(f"Error cargando template: {e}")
            raise

    def save_document(self, output_path: Path) -> bool:
        """Guardar documento generado"""
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(output_path)
            logger.info(f"Documento guardado: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error guardando documento: {e}")
            raise

    def replace_placeholder(self, placeholder: str, value: str) -> int:
        """Reemplazar placeholder en el documento"""
        replacements = 0

        for paragraph in self.document.paragraphs:
            if placeholder in paragraph.text:
                self._replace_text_in_paragraph(paragraph, placeholder, value)
                replacements += 1

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_text_in_paragraph(paragraph, placeholder, value)
                            replacements += 1

        return replacements

    @staticmethod
    def _replace_text_in_paragraph(paragraph, placeholder: str, value: str):
        """Reemplazar texto en un párrafo"""
        if placeholder in paragraph.text:
            new_text = paragraph.text.replace(placeholder, str(value))

            for run in paragraph.runs:
                run.text = ""

            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    def get_document(self) -> Document:
        """Obtener documento actual"""
        return self.document

    def clone_section(self, start_bookmark: str, end_bookmark: str, num_copies: int = 1):
        """Clonar una sección entre dos bookmarks"""
        logger.warning("clone_section no implementado aún en MVP")
        pass

    def get_placeholder_count(self) -> int:
        """Contar placeholders en el documento"""
        count = 0

        for paragraph in self.document.paragraphs:
            count += paragraph.text.count("{{")

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count += paragraph.text.count("{{")

        return count

